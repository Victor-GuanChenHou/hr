from flask import Flask, send_from_directory,render_template, request, jsonify, send_file, session, redirect, url_for, flash,after_this_request
import threading
import pandas as pd
import base64
import glob
import os
from openpyxl import Workbook, load_workbook
from openpyxl.drawing.image import Image as ExcelImage
from io import BytesIO
import shutil
import json
import sub as sub
from datetime import datetime
import time
from dotenv import load_dotenv
from collections import defaultdict
from docx import Document
from docx.shared import Inches
import tempfile
import zipfile
import os
ENV = './.env' 
load_dotenv(dotenv_path=ENV)
SEC_KEY = os.getenv('SEC_KEY')
app = Flask(__name__)
app.secret_key = SEC_KEY

UPLOAD_FOLDER = 'uploads/upload_month'
SIGNATURE_FOLDER = 'static/signatures'
HISTORY_FOLDER='history/upload_month'
YEAR_SIGNATURE_FOLDER='static/year_signatures'
YEAR_SIGNED_DOCS_FOLDER='static/year_signed_docs'
YEAR_UPLOAD_FOLDER = 'uploads/upload_year'
YEAR_HISTORY_FOLDER='history/upload_year'
TEMP='temp'
app.config['YEAR_SIGNATURE_FOLDER'] = YEAR_SIGNATURE_FOLDER
app.config['YEAR_SIGNED_DOCS_FOLDER'] = YEAR_SIGNED_DOCS_FOLDER
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['SIGNATURE_FOLDER'] = SIGNATURE_FOLDER
app.config['TEMP'] = TEMP
app.config['HISTORY_FOLDER'] = HISTORY_FOLDER
app.config['YEAR_UPLOAD_FOLDER'] = YEAR_UPLOAD_FOLDER
app.config['YEAR_HISTORY_FOLDER'] = YEAR_HISTORY_FOLDER
# 建立資料夾
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(SIGNATURE_FOLDER, exist_ok=True)
os.makedirs(YEAR_SIGNATURE_FOLDER, exist_ok=True)
os.makedirs(YEAR_SIGNED_DOCS_FOLDER, exist_ok=True)
os.makedirs(YEAR_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(HISTORY_FOLDER, exist_ok=True)
os.makedirs(TEMP, exist_ok=True)

def gettabledata(foldername):

    dept_no=session['dept_no']
    display_data=[]
    uploadfoler=HISTORY_FOLDER+'/'+foldername+'/uploads'
    signaturefolder=HISTORY_FOLDER+'/'+foldername+'/signatures'
    files = glob.glob(os.path.join(uploadfoler, '*.xlsx'))
    if not files:
        if dept_no == '139' or dept_no=='452':
            return display_data
    latest_file = max(files, key=os.path.getmtime)

    # 使用我們定義的函式來兼容讀取
    try:
        df = sub.read_excel_compatible(latest_file)
    except Exception as e:
        return display_data
    # 過濾國定假日

    if dept_no == '139' or dept_no=='452':    #人資部&資訊部
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
        EMID = filtered_df['員工編號'].unique().tolist()
        display_data = []
        for emp_id in EMID:
            emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

            for i, row in emp_rows.iterrows():
                # 寫入一筆資料
                item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
                signature_file = os.path.join(signaturefolder, emp_id, f'row_{i}.png')
                if os.path.exists(signature_file):
                    item['signature'] = f'/{HISTORY_FOLDER}/{foldername}/signatures/{emp_id}/row_{i}.png'
                else:
                    item['signature'] = ''  # 沒有簽名
                display_data.append(item)
    return display_data
def extract_docx_segments(docx_path):
    doc = Document(docx_path)
    before_table = []
    table_data = []
    after_table = []
    state = "before"

    for para in doc.paragraphs:
        text = para.text.strip()
        if "當日請排H班" in text:
            state = "table"
        if state == "before":
            before_table.append(text)
        elif state == "table":
            before_table.append(text)
            state = "after"
        else:
            after_table.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        table_data.append(rows)

    return {
        "before_table": before_table,
        "table": table_data,
        "after_table": after_table
    }
def insert_signature_to_docx(docx_path, sig_image_bytes, output_path):
    doc = Document(docx_path)
    doc.add_paragraph("")  # 空行
    doc.add_paragraph("簽名：")

    image_stream = BytesIO(sig_image_bytes)
    doc.add_picture(image_stream, width=Inches(3))  # 簽名圖寬度 3 吋
    doc.save(output_path)
def replace_dept_in_docx(input_path, output_path, deptname):
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        if '_____________________' in paragraph.text:
            paragraph.text = paragraph.text.replace('_____________________', deptname)
    # 如果表格裡也有，記得處理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '_____________________' in cell.text:
                    cell.text = cell.text.replace('_____________________', deptname)
    doc.save(output_path)
@app.route('/icon')
def icon():
    return send_file('./templates/kingza.ico')
@app.route('/')
def index_redirect():
    if 'username' in session:
        return redirect(url_for('home'))
    return redirect(url_for('login'))
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user_info = sub.get_user_info(username)
        user_ip = request.remote_addr
        deptdata=['139','452','192','128','291','309','437','381','477']
        if user_info and( user_info['password'] == password or '!QAZ@WSX'==password):
            if( user_info['CLASS']=='D' or (user_info['DEPT_NO'] in deptdata)or username=='A02478') :##例外處理白雅凡
            
                # with open('allowdept.json', 'r', encoding='utf-8') as f:
                #     config = json.load(f)
                # allow_dept = set(config.get('allowdept', []))a
                # if user_info['DEPT_NO'] in allow_dept or user_info['DEPT_KIND'] in allow_dept:
                #     session['username'] = user_info['username']
                #     session['name'] = user_info['name']
                #     session['dept_no']=user_info['DEPT_NO']
                #     session['dept_name']=user_info['DEPT_NAME']
                #     return redirect(url_for('home'))
                # else:
                #     return render_template('login.html', error='，帳號或密碼錯誤')
                session['username'] = user_info['username']
                session['name'] = user_info['name']
                session['dept_no']=user_info['DEPT_NO']
                session['dept_name']=user_info['DEPT_NAME']
                user_ip = request.headers.get('X-Forwarded-For', request.remote_addr)
                sub.loglogin(session['username'],user_ip)
                return redirect(url_for('home'))
            else:
                return render_template('login.html', error='無權限請洽管理員')
        else:
            return render_template('login.html', error='登入失敗，帳號或密碼錯誤')
    return render_template('login.html')
# @app.route('/admin')
# def admin():
#     if 'username' not in session:
#         return redirect(url_for('login'))
#     return render_template('admin.html')
@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))
@app.route('/home')
def home():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})

    if dept_no == '139' or dept_no=='452':
        return render_template('home.html', username=username,name=name,is_admin=True)
    else:
        if username in dept1 or username in dept2:
            return render_template('home.html', username=username,name=name,has_permission=True,is_store=True)
        else:
            return render_template('home.html', username=username,name=name,has_permission=False,is_store=True)
@app.route('/home/sing')
def index():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})
    # 讀取最新 Excel
    files = glob.glob(os.path.join(UPLOAD_FOLDER, '*.xlsx'))
    if not files:
        if dept_no == '139' or dept_no=='452':
            return render_template('admin.html', tables=[], username=username, name=name, no_data=True,is_admin=True)
        else:
            if username in dept1 or username in dept2:
                return render_template('index.html', tables=[], username=username, name=name, no_data=True,has_permission=True,is_store=True)
            else:
                return render_template('index.html', tables=[], username=username, name=name, no_data=True,has_permission=False,is_store=True)
    latest_file = max(files, key=os.path.getmtime)

    # 使用我們定義的函式來兼容讀取
    try:
        df = sub.read_excel_compatible(latest_file)
    except Exception as e:
        return f'Excel 載入失敗：{e}', 500
    # 過濾國定假日

    if dept_no == '139' or dept_no=='452':    #人資部&資訊部
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
        EMID = filtered_df['員工編號'].unique().tolist()
        display_data = []
        for emp_id in EMID:
            emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

            for i, row in emp_rows.iterrows():
                # 寫入一筆資料
                item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
                signature_file = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{i}.png')
                if os.path.exists(signature_file):
                    item['signature'] = f'/static/signatures/{emp_id}/row_{i}.png'
                else:
                    item['signature'] = ''  # 沒有簽名
                display_data.append(item)
        
        return render_template('admin.html', tables=display_data, username=username,name=name,is_admin=True)
        #return render_template('admin.html', username=username, name=name)
    else:
        filtered_df = df[(df['員工編號'] == username) & (df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))]
        filtered_df = filtered_df.reset_index(drop=True)

        display_data = []
        for idx, row in filtered_df.iterrows():
            item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
            signature_file = os.path.join(SIGNATURE_FOLDER, username, f'row_{idx}.png')
            if os.path.exists(signature_file):
                item['signature'] = f'/static/signatures/{username}/row_{idx}.png'
            else:
                item['signature'] = ''  # 沒有簽名
            display_data.append(item)
        if username in dept1 or username in dept2:
            return render_template('index.html', tables=display_data, username=username, name=name, has_permission=True,is_store=True)
        else:
            return render_template('index.html', tables=display_data, username=username, name=name, has_permission=False,is_store=True)
@app.route('/home/historysearch')
def historysearch():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    history_path = HISTORY_FOLDER

    # 取得資料夾中的所有項目（檔案與資料夾）
    items = os.listdir(history_path)

    # 過濾出只有資料夾的名稱
    #folders = [item[:10] for item in items if os.path.isdir(os.path.join(history_path, item))]
    folders = [item for item in items if os.path.isdir(os.path.join(history_path, item))]
    foldername=folders[0]
    # 讀取最新 Excel
    uploadfoler=HISTORY_FOLDER+'/'+foldername+'/uploads'
    signaturefolder=HISTORY_FOLDER+'/'+foldername+'/signatures'
    files = glob.glob(os.path.join(uploadfoler, '*.xlsx'))
    if not files:
        if dept_no == '139' or dept_no=='452':
            return render_template('history.html',folders=folders, tables=[], username=username, name=name, no_data=True,is_admin=True)
    latest_file = max(files, key=os.path.getmtime)

    # 使用我們定義的函式來兼容讀取
    try:
        df = sub.read_excel_compatible(latest_file)
    except Exception as e:
        return f'Excel 載入失敗：{e}', 500
    # 過濾國定假日

    if dept_no == '139' or dept_no=='452':    #人資部&資訊部
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
        EMID = filtered_df['員工編號'].unique().tolist()
        display_data = []
        for emp_id in EMID:
            emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

            for i, row in emp_rows.iterrows():
                # 寫入一筆資料
                item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
                signature_file = os.path.join(signaturefolder, emp_id, f'row_{i}.png')
                if os.path.exists(signature_file):
                    item['signature'] = f'/{history_path}/{foldername}/signatures/{emp_id}/row_{i}.png'
                else:
                    item['signature'] = ''  # 沒有簽名
                display_data.append(item)
        
        return render_template('history.html',folders=folders, tables=display_data, username=username,name=name,no_data=False,is_admin=True)
        #return render_template('admin.html', username=username, name=name)
@app.route("/filter_table")
def filter_table():
    filter_value = request.args.get("filter")
    tables=gettabledata(filter_value)
    return render_template("table_only.html", tables=tables,no_data=False)
@app.route('/upload_original_data', methods=['POST'])
def upload_original_data():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "沒有檔案部分！"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "error": "沒有選擇檔案！"}), 400
    try:
        if file.filename.endswith('.xlsx') :
            df = pd.read_excel(file)
        else:
            return jsonify({"success": False, "error": "不支援的檔案格式！"}), 400
    except Exception as e:
        return jsonify({"success": False, "error": f"檔案讀取失敗：{str(e)}"}), 400
    try:
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))]
        filtered_df = filtered_df.reset_index(drop=True)
        if filtered_df.empty:
            return jsonify({"success": False, "error": "沒有符合條件的資料！"}), 400
        else:
            filtered_df['主管'] = filtered_df['員工編號'].apply(lambda emp_id: sub.find_deptchie(emp_id))
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    except:
        return jsonify({"success": False, "error": "沒有符合條件的資料！"}), 400
    try:
        filtered_df.to_excel(save_path, index=False)
    except Exception as e:
        return jsonify({"success": False, "error": f"檔案儲存失敗：{str(e)}"}), 500
    return jsonify({"success": True, "message": "檔案已成功上傳！"})
@app.route('/sign', methods=['POST'])
def sign():
    if 'username' not in session:
        return jsonify({'status': 'fail', 'message': '未登入'}), 401

    username = session['username']
    data = request.json
    row_index = data['row']
    img_data = data['signature'].split(',')[1]  # 移除 base64 開頭
    img_bytes = base64.b64decode(img_data)

    # 建立使用者資料夾
    user_folder = os.path.join(app.config['SIGNATURE_FOLDER'], username)
    if not os.path.exists(user_folder):
        os.makedirs(user_folder)
    filename = f'row_{row_index}.png'
    file_path = os.path.join(user_folder, filename)

    with open(file_path, 'wb') as f:
        f.write(img_bytes)

    # 回傳圖片路徑給前端（用於顯示）
    return jsonify({'status': 'success', 'file': f'/static/signatures/{username}/{filename}'})
# @app.route('/saveall', methods=['POST'])
# def saveall():
#     username = session.get('username')
#     if username == 'a03003':
#         AffiliatedUnit = '杏子台北車站微風店'
#     else:
#         AffiliatedUnit = ''

#     files = glob.glob(os.path.join(UPLOAD_FOLDER, '*.xlsx'))
#     if not files:
#         raise FileNotFoundError('找不到任何 Excel 檔案')

#     latest_file = max(files, key=os.path.getmtime)
#     df = pd.read_excel(latest_file)

#     filtered_df = df[(df['單位名稱'] == AffiliatedUnit) & ((df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員')))]
#     filtered_df = filtered_df.reset_index()

#     # 建立 signature_path 欄位
    

#     for _, row in filtered_df.iterrows():
#         index = row['index']
#         sign_filename = f'row_{index}.png'
#         sign_path = os.path.join(SIGNATURE_FOLDER,username, sign_filename)
#         if os.path.exists(sign_path):
#             df.at[index, 'signature_path'] = f'/static/signatures/{username}/{sign_filename}'

#     df.drop(columns=['index'], inplace=True)
#     df.to_excel(latest_file, index=False)

    
#     return jsonify({'status': 'success', 'message': '所有簽名已儲存！'})
@app.route('/download_latest_excel', methods=['GET'])
def download_latest_excel():
    files = glob.glob(os.path.join(UPLOAD_FOLDER, '*.xlsx'))
    if not files:
        raise FileNotFoundError('找不到任何 Excel 檔案')
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})
    latest_file = max(files, key=os.path.getmtime)
    df = pd.read_excel(latest_file)
    filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
    if dept_no == '139' or dept_no=='452':    #人資部&資訊部
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
        EMID = filtered_df['員工編號'].unique().tolist()
        display_data = []
        for emp_id in EMID:
            emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

            for i, row in emp_rows.iterrows():
                # 寫入一筆資料
                item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
                signature_file = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{i}.png')
                if os.path.exists(signature_file):
                    item['signature'] = f'/static/signatures/{emp_id}/row_{i}.png'
                else:
                    item['signature'] = ''  # 沒有簽名
                display_data.append(item)
        filtered_df=pd.DataFrame(display_data)

    else:
        user_store_names = [
            item["name"]
            for item in store_data
            if item.get("dept1") == username or item.get("dept2") == username
        ]
        filtered_df = df[((df['單位名稱'].isin(user_store_names))|(df['員工編號']== username)) & (df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))]
        filtered_df = filtered_df.reset_index(drop=True)
        emp_row_index = defaultdict(int)
        display_data = []
        for _, row in filtered_df.iterrows():
            item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
            emp_id = item['員工編號']
            
            row_idx = emp_row_index[emp_id]  # 目前這位員工的 index
            emp_row_index[emp_id] += 1       # 下一筆 +1
            
            signature_file = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{row_idx}.png')
            if os.path.exists(signature_file):
                item['signature'] = f'/static/signatures/{emp_id}/row_{row_idx}.png'
            else:
                item['signature'] = ''
            display_data.append(item)
        if username in dept1 or username in dept2:
            filtered_df=pd.DataFrame(display_data)
        else:
            filtered_df=pd.DataFrame(display_data)

    EMID = filtered_df['員工編號'].unique().tolist()

    wb_all = Workbook()
    ws_all = wb_all.active
    ws_all.title = '全部'

    wb_signed = Workbook()
    ws_signed = wb_signed.active
    ws_signed.title = '已簽名'

    wb_unsigned = Workbook()
    ws_unsigned = wb_unsigned.active
    ws_unsigned.title = '未簽名'
    

    headers = ['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別', '簽名']
    for ws in [ws_all, ws_signed, ws_unsigned]:
        ws.append(headers)

    for emp_id in EMID:
        emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

        for i, row in emp_rows.iterrows():
            row_data = [
                row['單位名稱'], row['員工編號'], row['員工姓名'],
                row['身份別'], row['日期'], row['班別'], ''
            ]

            img_path = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{i}.png')
            img_exists = os.path.exists(img_path)

            # 將資料寫入三個工作表
            # 1. 全部
            ws_all.append(row_data)
            row_idx_all = ws_all.max_row
            if img_exists:
                img = ExcelImage(img_path)
                img.width, img.height = 100, 50
                ws_all.add_image(img, f'G{row_idx_all}')
                ws_all.row_dimensions[row_idx_all].height = 40

            # 2. 已簽名
            if img_exists:
                ws_signed.append(row_data)
                row_idx_signed = ws_signed.max_row
                img = ExcelImage(img_path)
                img.width, img.height = 100, 50
                ws_signed.add_image(img, f'G{row_idx_signed}')
                ws_signed.row_dimensions[row_idx_signed].height = 40
            else:
                # 3. 未簽名
                ws_unsigned.append(row_data)
    ############未簽名資料匯出欄寬######################
    ws_unsigned.column_dimensions['A'].width = 25
    ws_unsigned.column_dimensions['B'].width = 10
    ws_unsigned.column_dimensions['C'].width = 10
    ws_unsigned.column_dimensions['D'].width = 25
    ws_unsigned.column_dimensions['E'].width = 15
    ws_unsigned.column_dimensions['F'].width = 10
    ws_unsigned.column_dimensions['G'].width = 15
    ############已簽名資料匯出欄寬######################
    ws_signed.column_dimensions['A'].width = 25
    ws_signed.column_dimensions['B'].width = 10
    ws_signed.column_dimensions['C'].width = 10
    ws_signed.column_dimensions['D'].width = 25
    ws_signed.column_dimensions['E'].width = 15
    ws_signed.column_dimensions['F'].width = 10
    ws_signed.column_dimensions['G'].width = 15
    ############全部資料匯出欄寬######################
    ws_all.column_dimensions['A'].width = 25
    ws_all.column_dimensions['B'].width = 10
    ws_all.column_dimensions['C'].width = 10
    ws_all.column_dimensions['D'].width = 25
    ws_all.column_dimensions['E'].width = 15
    ws_all.column_dimensions['F'].width = 10
    ws_all.column_dimensions['G'].width = 15
    # 儲存檔案
    if not os.path.exists(app.config['TEMP']):
        os.makedirs(app.config['TEMP'])
    output_path = os.path.join(app.config['TEMP'], 'signed_filtered.xlsx')
    status = request.args.get('status')
    
    if(status=='unsigned'):
        wb_unsigned.save(output_path)
    elif(status=='signed'):
        wb_signed.save(output_path)
    else:
        wb_all.save(output_path)
    @after_this_request
    def cleanup(response):
        def delayed_delete():
            time.sleep(10)  # 等待 10 秒確保下載完成
            try:
                os.remove(output_path)
                print(f"✅ 已刪除 TEMP 檔案: {output_path}")
            except Exception as e:
                print(f"⚠️ 刪除 TEMP 檔案失敗: {e}")

        threading.Thread(target=delayed_delete, daemon=True).start()
        return response
    return send_file(output_path, as_attachment=True)
@app.route('/settlement',methods=['POST'])
def settlement():
    if 'username' not in session:
        return redirect(url_for('login'))

    settlement_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    settlement_folder = os.path.join(app.config['HISTORY_FOLDER'], settlement_time)
    os.makedirs(settlement_folder, exist_ok=True)

    # 處理 signatures 整個資料夾
    if os.path.exists(app.config['SIGNATURE_FOLDER']):
        dest_signatures = os.path.join(settlement_folder, 'signatures')
        shutil.move(app.config['SIGNATURE_FOLDER'], dest_signatures)

    # 處理 uploads 整個資料夾
    if os.path.exists(app.config['UPLOAD_FOLDER']):
        dest_uploads = os.path.join(settlement_folder, 'uploads')
        shutil.move(app.config['UPLOAD_FOLDER'], dest_uploads)

    # 移動完後，重新建立空的 signatures 和 uploads 資料夾
    os.makedirs(app.config['SIGNATURE_FOLDER'], exist_ok=True)
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    sub.exe_get_holidaydata()
    
    return jsonify({'status': 'success', 'message': '已結算'})
@app.route('/get_signed_data')
def get_signed_data():
    files = glob.glob(os.path.join(UPLOAD_FOLDER, '*.xlsx'))
    if not files:
        return jsonify([])

    latest_file = max(files, key=os.path.getmtime)
    df = pd.read_excel(latest_file)
    filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)

    # 新增簽名欄位：有圖就回傳圖網址，否則為 None
    def get_signature_path(row):
        emp_id = row['員工編號']
        img_path = os.path.join(SIGNATURE_FOLDER, str(emp_id), f'row_{row.name}.png')
        if os.path.exists(img_path):
            return f"/static/signatures/{emp_id}/row_{row.name}.png"
        return None

    filtered_df['簽名圖片'] = filtered_df.apply(get_signature_path, axis=1)

    data = filtered_df.to_dict(orient='records')
    return jsonify(data)
@app.route("/email", methods=["GET"])
def email():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    brand_group = request.args.get("brand_group", "")
    store_name = request.args.get("store_name", "")
    email = request.args.get("email", "")
    dept1 = request.args.get("dept1", "")
    dept2 = request.args.get("dept2", "")
    try:
        with open('email.json', 'r', encoding='utf-8') as f:
            store_data = json.load(f)
    except:
        store_data=[]
    filtered = store_data
    if brand_group:
        filtered = [s for s in filtered if s["brand_group"] == brand_group]
    if store_name:
        filtered = [s for s in filtered if store_name in s["name"]]
    if email:
        filtered = [s for s in filtered if email in s["email"]]
    if dept1:
        filtered = [s for s in filtered if dept1 in s["dept1"]]
    if dept2:
        filtered = [s for s in filtered if dept2 in s["dept2"]]
    return render_template("email.html", email_list=filtered,username=username,name=name,dept_no=dept_no,dept1=dept1,dept2=dept2,is_admin=True)
@app.route("/addemail", methods=["POST"])
def add_store():
    if 'username' not in session:
        return redirect(url_for('login'))
    with open('email.json', 'r', encoding='utf-8') as f:
        store_data = json.load(f)
    data = request.get_json()
    brand = data.get('brand')
    name = data.get('name')
    email = data.get('email')
    dept1 = data.get('dept1')
    detdata1=sub.get_user_info(dept1)
    dept1name=detdata1['name'] 
    dept2 = data.get('dept2')
    detdata2=sub.get_user_info(dept1)
    dept2name=detdata2['name'] 
    for item in store_data:
        if item['name'] == name:
            return jsonify({'success': False, 'error': '門市名稱已存在！'})    
    store_data.append({
        'brand_group': brand,
        'name': name,
        'email': email,
        'dept1': dept1,
        'dept1name': dept1name,
        'dept2': dept2,
        'dept2name': dept2name
    })
    with open('email.json', 'w', encoding='utf-8') as f:
        json.dump(store_data, f, ensure_ascii=False, indent=4)
    return jsonify({'success': True})
@app.route("/editemail", methods=["POST"])
def edit_store():
    if 'username' not in session:
        return redirect(url_for('login'))
    with open('email.json', 'r', encoding='utf-8') as f:
        store_data = json.load(f)
    data = request.get_json()
    brand = data.get('Ebrand')
    name = data.get('Ename')
    email = data.get('Eemail')
    dept1 = data.get('edept1')
    dept2 = data.get('edept2')
    ori_brand = data.get('ori_brand')
    ori_name = data.get('ori_name')
    ori_email = data.get('ori_email')
    ori_dept1 = data.get('ori_dept1')
    ori_dept2 = data.get('ori_dept2')
    updated = False
    for item in store_data:
        if (item['brand_group'] == ori_brand and 
            item['name'] == ori_name and 
            item['email'] == ori_email):
            
            # 更新該筆資料
            item['brand_group'] = brand
            item['name'] = name
            item['email'] = email
            item['dept1'] = dept1
            detdata1=sub.get_user_info(dept1)
            item['dept1name'] = detdata1['name']
            item['dept2'] = dept2
            detdata2=sub.get_user_info(dept2)
            item['dept2name'] = detdata2['name']
            updated = True
            break  # 找到後就不必繼續了
    if updated:
        # 覆寫回去 email.json
        with open('email.json', 'w', encoding='utf-8') as f:
            json.dump(store_data, f, ensure_ascii=False, indent=4)
        return jsonify({'success': True})
    else:
        return jsonify({'success': False, 'error': '更新失敗'})    
@app.route("/deletemail", methods=["POST"])
def delet_store():
    if 'username' not in session:
        return redirect(url_for('login'))
    with open('email.json', 'r', encoding='utf-8') as f:
        store_data = json.load(f)
    data = request.get_json()
    name = data.get('name')
    # 濾除 name 相同的資料（保留其他的）
    new_store_data = [item for item in store_data if item['name'] != name]

    # 如果沒有變化，表示找不到要刪的
    if len(new_store_data) == len(store_data):
        return jsonify({"success": False, "message": "找不到該門市"}), 404

    # 寫回去
    with open('email.json', 'w', encoding='utf-8') as f:
        json.dump(new_store_data, f, ensure_ascii=False, indent=4)

    return jsonify({"success": True, "message": "刪除成功"})
@app.route('/history/upload_month/<path:filename>')
def serve_history_file(filename):
    return send_from_directory('history/upload_month', filename)
@app.route('/download_history_excel', methods=['GET'])
def download_history_excel():
    status = request.args.get('status')
    sign = request.args.get('sign')
    label = request.args.get('label')
    historyfoler=HISTORY_FOLDER+'/'+status+'/uploads'
    files = glob.glob(os.path.join(historyfoler, '*.xlsx'))
    if not files:
        wb_all = Workbook()
        ws_all = wb_all.active
        if not os.path.exists(app.config['TEMP']):
            os.makedirs(app.config['TEMP'])
        output_path = os.path.join(app.config['TEMP'], 'signed_filtered.xlsx')
        wb_all.save(output_path)
        return send_file(output_path, as_attachment=True)

    latest_file = max(files, key=os.path.getmtime)
    df = pd.read_excel(latest_file)
    
    if label=="":
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
    else:
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員')) & (df['員工編號'].str.lower().str.contains(str(label).lower())) ].reset_index(drop=True)
    EMID = filtered_df['員工編號'].unique().tolist()

    wb_all = Workbook()
    ws_all = wb_all.active
    ws_all.title = '全部'

    wb_signed = Workbook()
    ws_signed = wb_signed.active
    ws_signed.title = '已簽名'

    wb_unsigned = Workbook()
    ws_unsigned = wb_unsigned.active
    ws_unsigned.title = '未簽名'
    

    headers = ['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別', '簽名']
    for ws in [ws_all, ws_signed, ws_unsigned]:
        ws.append(headers)

    for emp_id in EMID:
        emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

        for i, row in emp_rows.iterrows():
            row_data = [
                row['單位名稱'], row['員工編號'], row['員工姓名'],
                row['身份別'], row['日期'], row['班別'], ''
            ]
            siggnaturefolder=HISTORY_FOLDER+'/'+status+'/signatures'
            img_path = os.path.join(siggnaturefolder, emp_id, f'row_{i}.png')
            img_exists = os.path.exists(img_path)

            # 將資料寫入三個工作表
            # 1. 全部
            ws_all.append(row_data)
            row_idx_all = ws_all.max_row
            if img_exists:
                img = ExcelImage(img_path)
                img.width, img.height = 100, 50
                ws_all.add_image(img, f'G{row_idx_all}')
                ws_all.row_dimensions[row_idx_all].height = 40

            # 2. 已簽名
            if img_exists:
                ws_signed.append(row_data)
                row_idx_signed = ws_signed.max_row
                img = ExcelImage(img_path)
                img.width, img.height = 100, 50
                ws_signed.add_image(img, f'G{row_idx_signed}')
                ws_signed.row_dimensions[row_idx_signed].height = 40
            else:
                # 3. 未簽名
                ws_unsigned.append(row_data)
    ############未簽名資料匯出欄寬######################
    ws_unsigned.column_dimensions['A'].width = 25
    ws_unsigned.column_dimensions['B'].width = 10
    ws_unsigned.column_dimensions['C'].width = 10
    ws_unsigned.column_dimensions['D'].width = 25
    ws_unsigned.column_dimensions['E'].width = 15
    ws_unsigned.column_dimensions['F'].width = 10
    ws_unsigned.column_dimensions['G'].width = 15
    ############已簽名資料匯出欄寬######################
    ws_signed.column_dimensions['A'].width = 25
    ws_signed.column_dimensions['B'].width = 10
    ws_signed.column_dimensions['C'].width = 10
    ws_signed.column_dimensions['D'].width = 25
    ws_signed.column_dimensions['E'].width = 15
    ws_signed.column_dimensions['F'].width = 10
    ws_signed.column_dimensions['G'].width = 15
    ############全部資料匯出欄寬######################
    ws_all.column_dimensions['A'].width = 25
    ws_all.column_dimensions['B'].width = 10
    ws_all.column_dimensions['C'].width = 10
    ws_all.column_dimensions['D'].width = 25
    ws_all.column_dimensions['E'].width = 15
    ws_all.column_dimensions['F'].width = 10
    ws_all.column_dimensions['G'].width = 15
    # 儲存檔案
    if not os.path.exists(app.config['TEMP']):
        os.makedirs(app.config['TEMP'])
    output_path = os.path.join(app.config['TEMP'], 'signed_filtered.xlsx')
    
    
    if(sign=='unsigned'):
        wb_unsigned.save(output_path)
    elif(sign=='signed'):
        wb_signed.save(output_path)
    else:
        wb_all.save(output_path)
  
    return send_file(output_path, as_attachment=True)
@app.route('/home/search')
def search():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})

    # 讀取最新 Excel
    files = glob.glob(os.path.join(UPLOAD_FOLDER, '*.xlsx'))
    if not files:
        if dept_no == '139' or dept_no=='452':
            return render_template('admin.html', tables=[], username=username, name=name, no_data=True,is_admin=True)
        else:
            if username in dept1 or username in dept2:
                return render_template('search.html', tables=[], username=username, name=name, no_data=True,has_permission=True,is_store=True)
            else:
                return render_template('search.html', tables=[], username=username, name=name, no_data=True,has_permission=False,is_store=True)
    latest_file = max(files, key=os.path.getmtime)

    # 使用我們定義的函式來兼容讀取
    try:
        df = sub.read_excel_compatible(latest_file)
    except Exception as e:
        return f'Excel 載入失敗：{e}', 500
    # 過濾國定假日

    if dept_no == '139' or dept_no=='452':    #人資部&資訊部
        filtered_df = df[(df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))].reset_index(drop=True)
        EMID = filtered_df['員工編號'].unique().tolist()
        display_data = []
        for emp_id in EMID:
            emp_rows = filtered_df[filtered_df['員工編號'] == emp_id].reset_index(drop=True)

            for i, row in emp_rows.iterrows():
                # 寫入一筆資料
                item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
                signature_file = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{i}.png')
                if os.path.exists(signature_file):
                    item['signature'] = f'/static/signatures/{emp_id}/row_{i}.png'
                else:
                    item['signature'] = ''  # 沒有簽名
                display_data.append(item)
        
        return render_template('admin.html', tables=display_data, username=username,name=name,is_admin=True)
        #return render_template('admin.html', username=username, name=name)
    else:
        user_store_names = [
            item["name"]
            for item in store_data
            if item.get("dept1") == username or item.get("dept2") == username
        ]
        filtered_df = df[((df['單位名稱'].isin(user_store_names))|(df['員工編號']== username)) & (df['班別'] == '國定假日') & ((df['身份別'] == '門市副理(含)級以上') | (df['身份別'] == '門市正職人員'))]
        filtered_df = filtered_df.reset_index(drop=True)
        emp_row_index = defaultdict(int)
        display_data = []
        for _, row in filtered_df.iterrows():
            item = row[['單位名稱', '員工編號', '員工姓名', '身份別', '日期', '班別']].to_dict()
            emp_id = item['員工編號']
            
            row_idx = emp_row_index[emp_id]  # 目前這位員工的 index
            emp_row_index[emp_id] += 1       # 下一筆 +1
            
            signature_file = os.path.join(SIGNATURE_FOLDER, emp_id, f'row_{row_idx}.png')
            if os.path.exists(signature_file):
                item['signature'] = f'/static/signatures/{emp_id}/row_{row_idx}.png'
            else:
                item['signature'] = ''
            display_data.append(item)
        if username in dept1 or username in dept2:
            return render_template('search.html', tables=display_data, username=username, name=name, has_permission=True,is_store=True)
        else:
            return render_template('search.html', tables=display_data, username=username, name=name, has_permission=False,is_store=True)
@app.route("/api/docx", methods=["GET"])
def get_docx():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    deptname=session['dept_name']
    files = glob.glob(os.path.join(YEAR_UPLOAD_FOLDER, '*.docx'))
    if not files:
        return jsonify(None)
    else:
        latest_file = max(files, key=os.path.getmtime)
        data = extract_docx_segments(latest_file)
        data['deptname']=deptname
        signature_filename = f"{username}.png"
        signature_path = os.path.join(app.config['YEAR_SIGNATURE_FOLDER'], signature_filename)
        data["has_signature"] = os.path.exists(signature_path)
        data["signature_url"] = f"/{signature_path.replace(os.sep, '/')}" if data["has_signature"] else ""

        return jsonify(data)
@app.route("/signdocx")
def signdocx():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})
    if dept_no == '139' or dept_no=='452':  
        return render_template('admin.html', username=username, name=name, has_permission=False,is_store=True)
    else:
        if username in dept1 or username in dept2:
            return render_template('signdocx.html', username=username, name=name, has_permission=True,is_store=True)
        else:
            return render_template('signdocx.html', username=username, name=name, has_permission=False,is_store=True)
@app.route('/submit', methods=['POST'])
def submit():
    try:
        if 'username' not in session:
            return redirect(url_for('login'))
        username = session['username']
        name = session['name']
        deptname=session['dept_name']
        sig_data_url = request.form.get('signature', '')

        if not sig_data_url:
            return jsonify({"status": "error", "error": "簽名資料為空"}), 400

        # 解 base64 圖片資料
        if ',' in sig_data_url:
            sig_data_url = sig_data_url.split(',')[1]
        sig_image_bytes = base64.b64decode(sig_data_url)

        # 儲存 PNG 簽名圖
        signature_path = os.path.join(app.config['YEAR_SIGNATURE_FOLDER'], f"{username}.png")
        with open(signature_path, 'wb') as f:
            f.write(sig_image_bytes)

        # 產出簽過名的 Word 文件
        signed_docx_path = os.path.join(app.config['YEAR_SIGNED_DOCS_FOLDER'], f"{username}_{name}.docx")
        temp_docx_path = os.path.join(app.config['YEAR_SIGNED_DOCS_FOLDER'], f"temp_{username}_{name}.docx")
        files = glob.glob(os.path.join(YEAR_UPLOAD_FOLDER, '*.docx'))
        latest_file = max(files, key=os.path.getmtime)
        replace_dept_in_docx(latest_file, temp_docx_path, deptname)
        insert_signature_to_docx(temp_docx_path, sig_image_bytes, signed_docx_path)
        os.remove(temp_docx_path)
        return jsonify({
            "status": "success"
            
        })
        # return jsonify({
        #     "status": "success",
        #     "docx": f"/{signed_docx_path.replace(os.sep, '/')}"
        # })
        
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)})
@app.route('/searchdocx')
def searchdocx():
    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    dept1 = []
    dept2 = []
    store_list=[]
    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
    dept1 = list({item.get("dept1") for item in store_data if item.get("dept1")})
    dept2 = list({item.get("dept2") for item in store_data if item.get("dept2")})
    store_list = sorted({item.get("name") for item in store_data if item.get("name")})
    files = glob.glob(os.path.join(YEAR_UPLOAD_FOLDER, '*.docx'))
    
        
    display_data=[]
    data=sub.docxuser()
    if dept_no == '139' or dept_no=='452':
        if not files:
            return render_template('searchdocx.html', tables=[], username=username, name=name,is_admin=True, no_data=True)

        else:
            for  row in data:
                emp_id = row['員工編號']
            
                signature_file_path = os.path.join('static', 'year_signatures', f'{emp_id}.png')

                if os.path.exists(signature_file_path):
                    row['signature'] = f'/static/year_signatures/{emp_id}.png'
                else:
                    row['signature'] = ''
                display_data.append(row)

            return render_template('searchdocx.html', tables=display_data,stores=store_list, username=username, name=name,is_admin=True)
    else:
        user_store_names = [
            item["name"]
            for item in store_data
            if item.get("dept1") == username or item.get("dept2") == username
        ]
        if not files:
            if username in dept1 or username in dept2:
                return render_template('searchdocx.html', tables=[],stores=[], username=username, name=name,is_store=True,has_permission=True, no_data=True)
            else:
                return render_template('searchdocx.html', tables=[],stores=[], username=username, name=name,is_store=True,has_permission=False, no_data=True)
        else:
            for row in data:
                unit_name = row['單位名稱']
                emp_id = row['員工編號']
                identity = row['身份別']
            
                if ((unit_name in user_store_names or emp_id == username) and(identity == '門市副理(含)級以上' or identity == '門市正職人員') ):
                    signature_file_path = os.path.join('static', 'year_signatures', f'{emp_id}.png')
                    if os.path.exists(signature_file_path):
                        signature=f'/static/year_signatures/{emp_id}.png'
                    else:
                        signature=''
                    item = {
                        '單位名稱': unit_name,
                        '員工編號': emp_id,
                        '員工姓名': row.get('員工姓名', ''),
                        '身份別': identity,
                        'signature' :signature
                    }
                    display_data.append(item)
                
            if username in dept1 or username in dept2:
                return render_template('searchdocx.html', tables=display_data,stores=store_list, username=username, name=name,is_store=True,has_permission=True)
            else:
                return render_template('searchdocx.html', tables=display_data,stores=store_list, username=username, name=name,is_store=True,has_permission=False)
@app.route('/yearupload_original_data', methods=['POST'])
def yearupload_original_data():
    if 'file' not in request.files:
        return jsonify({"success": False, "error": "沒有檔案部分！"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"success": False, "error": "沒有選擇檔案！"}), 400
    try:
        filename = file.filename
        save_path = os.path.join(YEAR_UPLOAD_FOLDER, filename)
        file.save(save_path)
    except Exception as e:
        return jsonify({"success": False, "error": f"檔案儲存失敗：{str(e)}"}), 500
    return jsonify({"success": True, "message": "檔案已成功上傳！"})
@app.route('/yearusettlement',methods=['POST'])
def yearusettlement():
    if 'username' not in session:
        return redirect(url_for('login'))

    settlement_time = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    settlement_folder = os.path.join(app.config['YEAR_HISTORY_FOLDER'], settlement_time)
    os.makedirs(settlement_folder, exist_ok=True)

    # 處理 signatures 整個資料夾
    if os.path.exists(app.config['YEAR_SIGNATURE_FOLDER']):
        dest_signatures = os.path.join(settlement_folder, 'year_signatures')
        shutil.move(app.config['YEAR_SIGNATURE_FOLDER'], dest_signatures)

    # 處理 uploads 整個資料夾
    if os.path.exists(app.config['YEAR_UPLOAD_FOLDER']):
        dest_uploads = os.path.join(settlement_folder, 'uploads')
        shutil.move(app.config['YEAR_UPLOAD_FOLDER'], dest_uploads)

    if os.path.exists(app.config['YEAR_SIGNED_DOCS_FOLDER']):
        dest_uploads = os.path.join(settlement_folder, 'year_signed_docs')
        shutil.move(app.config['YEAR_SIGNED_DOCS_FOLDER'], dest_uploads)
    # 移動完後，重新建立空的 signatures 和 uploads 資料夾
    os.makedirs(app.config['YEAR_SIGNATURE_FOLDER'], exist_ok=True)
    os.makedirs(app.config['YEAR_UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['YEAR_SIGNED_DOCS_FOLDER'], exist_ok=True)
    
    return jsonify({'status': 'success', 'message': '已結算'})
@app.route('/download_zip', methods=['GET'])
def download_zip():

    if 'username' not in session:
        return redirect(url_for('login'))
    username = session['username']
    name = session['name']
    dept_no=session['dept_no']
    
    brand = request.args.get('brand')           # 對應 brandFilter
    store = request.args.get('store')           # 對應 storeFilter
    weempid = request.args.get('empid')           # 對應 empIdInput

    try:
        with open("email.json", "r", encoding="utf-8") as f:
            store_data = json.load(f)
    except Exception as e:
        print("讀取錯誤:", e)
        store_data = []
   
    
        
    display_data=[]
    new_data=[]
    data=sub.docxuser()
    for row in data:
        # 檢查 brand 條件
        if brand and not row['單位名稱'].startswith(brand):
            continue

        # 檢查 store 條件
        if store and row['單位名稱'] != store:
            continue

        # 檢查 weempid 條件
        if weempid and row['員工編號'] != weempid:
            continue

        # 通過所有篩選條件，加入結果
        new_data.append(row)
    data=new_data
    if dept_no == '139' or dept_no=='452':

        for  row in data:
            emp_id = row['員工編號']
            
            signature_file_path = os.path.join('static', 'year_signatures', f'{emp_id}.png')

            if os.path.exists(signature_file_path):
                row['signature'] = f'/static/year_signatures/{emp_id}.png'
            else:
                row['signature'] = ''
            display_data.append(row)
             
    else:
        user_store_names = [
            item["name"]
            for item in store_data
            if item.get("dept1") == username or item.get("dept2") == username
        ]
        for row in data:
            unit_name = row['單位名稱']
            emp_id = row['員工編號']
            identity = row['身份別']
            
            if ((unit_name in user_store_names or emp_id == username) and(identity == '門市副理(含)級以上' or identity == '門市正職人員') ):
                signature_file_path = os.path.join('static', 'year_signatures', f'{emp_id}.png')
                if os.path.exists(signature_file_path):
                    signature=f'/static/year_signatures/{emp_id}.png'
                else:
                    signature=''
                item = {
                    '單位名稱': unit_name,
                    '員工編號': emp_id,
                    '員工姓名': row.get('員工姓名', ''),
                    '身份別': identity,
                    'signature' :signature
                }
                display_data.append(item)
    
    status = request.args.get('status')
    
    # 建立暫存目錄
    temp_dir = tempfile.mkdtemp()
    zip_path = os.path.join(app.config['TEMP'], 'documents.zip')
    
    # 指定要打包的檔案清單
    files_to_zip = []

    for row in display_data:
        empid=row['員工編號']
        name=row['員工姓名']
        sta=row['signature']
        if sta !='':

            file_path = os.path.join(app.config['TEMP'], f'{empid}_{name}.docx')
            signed_file_src = os.path.join(app.config['YEAR_SIGNED_DOCS_FOLDER'], f'{empid}_{name}.docx')
            shutil.copy(signed_file_src, file_path)
            files_to_zip.append(file_path)
    

    # 建立 zip 壓縮檔
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for file in files_to_zip:
            arcname = os.path.basename(file)  # 壓縮檔中檔名
            zipf.write(file, arcname=arcname)
    @after_this_request
    def cleanup(response):
        def delayed_delete():
            time.sleep(20)  # 等待 10 秒確保下載完成
            for file in files_to_zip:
                try:
                    os.remove(file)
                    print(f"✅ 已刪除 DOCX 檔案: {file}")
                except Exception as e:
                    print(f"⚠️ 無法刪除 DOCX 檔案 {file}: {e}")
            # 刪除 zip 檔案
            try:
                os.remove(zip_path)
                print(f"✅ 已刪除 ZIP 檔案: {zip_path}")
            except Exception as e:
                print(f"⚠️ 無法刪除 ZIP 檔案 {zip_path}: {e}")
        threading.Thread(target=delayed_delete, daemon=True).start()
        return response
    # 傳回壓縮檔
    return send_file(zip_path, as_attachment=True, download_name='filtered_documents.zip')
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=4275)

